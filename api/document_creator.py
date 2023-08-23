import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.shared import Inches
import os


def create_docx(ebook, ebook_id):
    document = docx.Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    title = document.add_heading(ebook['title'], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if "cover" in ebook.keys():
        image = document.add_picture(os.getcwd() + f"//output/ebook{ebook_id}.png", Inches(7))
        image.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_page_break()
    tbo = document.add_heading('Table of Contents', 1)
    tbo.bold = True
    for chapter in ebook['table_of_contents']:
        document.add_paragraph(chapter)
    document.add_page_break()
    for chapter in ebook['chapters'].keys():
        title_c = document.add_heading(ebook['chapters'][chapter]['title'], 1)
        title_c.bold = True
        title_c.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(ebook['chapters'][chapter]['text'])
        document.add_page_break()
    document.add_heading('SOURCES', 1)
    for s in ebook['sources']:
        document.add_paragraph(s)
    document.save(f"output//ebook{ebook_id}.docx")


def convert_docx_to_pdf(ebook, ebook_id):
    create_docx(ebook, ebook_id)
    convert(f"/output/ebook{ebook_id}.docx", f"/output/ebook{ebook_id}.pdf")
